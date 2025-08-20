import os
import re
import zipfile
import io
import jinja2
from bs4 import BeautifulSoup
from docxtpl import DocxTemplate, RichText, InlineImage
from docx.shared import Inches
from models import report as report_model, template as template_model
import datetime
from lxml import etree
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GENERATED_DOCS_PATH = "storage/generated"

def _add_toc(doc):
    """
    在文档的开头（第一个段落之后）插入一个目录域。
    """
    # 创建一个新的段落用于放置目录
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    # 创建 fldChar begin 元素
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    
    # 创建 instrText 元素 (目录指令)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u' # 标准的目录指令
    
    # 创建 fldChar separate 元素
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    
    # 创建 fldChar end 元素
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    # 将所有元素添加到 run 中
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    run._r.append(fldChar_end)
    
    # 将新创建的目录段落移动到文档的开头（标题之后）
    # 假设文档的第一个元素是标题，我们将目录放在第二个位置
    if len(doc.element.body) > 1:
        doc.element.body.insert(1, paragraph._p)


def _enable_auto_update_fields(docx_path):
    """
    通过修改 docx 文件包中的 settings.xml，强制 Word 在打开时更新域（如目录）。
    """
    namespace = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }
    update_fields_tag = '{' + namespace['w'] + '}updateFields'
    val_attribute = '{' + namespace['w'] + '}val'

    # 使用 zipfile 库直接操作 docx (它是一个 zip 文件)
    zin = zipfile.ZipFile(docx_path, 'r')
    zout = zipfile.ZipFile(docx_path, 'a')

    # 解压 settings.xml
    xml_content = zin.read('word/settings.xml')
    
    # 使用 lxml 解析 XML
    root = etree.fromstring(xml_content)
    
    # 检查 <w:updateFields> 是否已存在
    update_fields_element = root.find('w:updateFields', namespaces=namespace)
    if update_fields_element is None:
        # 如果不存在，则创建一个新的元素并添加到 settings 根元素的末尾
        update_fields_element = etree.Element(update_fields_tag, nsmap=namespace)
        root.append(update_fields_element)
    
    # 设置 w:val="true"
    update_fields_element.set(val_attribute, "true")

    # 将修改后的 XML 写回 zip 文件
    # 注意：lxml 自动处理命名空间前缀，所以我们不需要手动添加 'w:'
    xml_str = etree.tostring(root, pretty_print=True)
    zout.writestr('word/settings.xml', xml_str)
    
    zin.close()
    zout.close()


def format_date(date_string):
    """将ISO格式的日期字符串转换为 'YYYY年M月D日' 格式"""
    if not date_string:
        return ''
    try:
        # 解析ISO 8601格式的日期时间字符串
        dt = datetime.datetime.fromisoformat(date_string.replace('Z', '+00:00'))
        # 格式化为 "YYYY年M月D日"
        return dt.strftime("%Y年%m月%d日")
    except (ValueError, TypeError):
        # 如果格式不正确或不是字符串，返回原始值
        return date_string

def render_docx(report: report_model.Report, template: template_model.Template) -> str:
    if not os.path.exists(template.filepath):
        raise FileNotFoundError("模板文件不存在")

    doc = DocxTemplate(template.filepath)
    
    image_placeholders = {}
    image_counter = 0

    def process_html_for_docx(html_string):
        nonlocal image_counter
        if not html_string: return ''
        soup = BeautifulSoup(html_string, 'lxml')
        for img in soup.find_all('img'):
            src = img.get('src')
            if src:
                local_path = src.lstrip('/')
                if os.path.exists(local_path):
                    placeholder = f"##IMG_{image_counter}##"
                    image_placeholders[placeholder] = local_path
                    img.replace_with(placeholder)
                    image_counter += 1
        return soup.get_text(separator='\n')

    vuls = [{'vul_name': v.vul_name or '', 'vul_level': {'High': '高危', 'Medium': '中危', 'Low': '低危'}.get(v.vul_level, '未知'), 'vul_describe': v.vul_describe or '', 'vul_url': v.vul_url or '', 'vul_analysis': process_html_for_docx(v.vul_analysis), 'vul_modify_repair': v.vul_modify_repair or ''} for v in report.vuls]

    context = {
        'report_center': report.report_center or '', 'report_systemname': report.report_systemname or '',
        'report_start_time': format_date(report.report_start_time), 'report_end_time': format_date(report.report_end_time),
        'author': report.author or '', 'reviewer': report.reviewer or '', 'report_center_short': report.report_center_short or '',
        'overall_risk_level': report.overall_risk_level or '', 'vuls': vuls, 'vuls_length': len(vuls),
        'vul_stats': {'high': sum(1 for v in vuls if v.get('vul_level') == '高危'), 'medium': sum(1 for v in vuls if v.get('vul_level') == '中危'), 'low': sum(1 for v in vuls if v.get('vul_level') == '低危')},
        'targets': [{'name': t.name, 'url': t.url} for t in report.targets],
        'members': [{'role': m.role, 'name': m.name, 'contact': m.contact} for m in report.members],
        'first_vul_name': vuls[0]['vul_name'] if vuls else '',
    }

    doc.render(context)

    # --- 图片替换 ---
    for p in doc.paragraphs:
        if any('##IMG_' in run.text for run in p.runs):
            inline_runs = []
            for run in p.runs:
                if '##IMG_' in run.text:
                    parts = re.split(r'(##IMG_\d+##)', run.text)
                    for part in parts:
                        if part in image_placeholders:
                            inline_runs.append({'type': 'image', 'path': image_placeholders[part]})
                        elif part:
                            inline_runs.append({'type': 'text', 'content': part})
                else:
                    inline_runs.append({'type': 'text', 'content': run.text})
            
            p.clear()
            for item in inline_runs:
                run = p.add_run()
                if item['type'] == 'text':
                    run.text = item['content']
                elif item['type'] == 'image':
                    try:
                        run.add_picture(item['path'], width=Inches(5.5))
                    except Exception as e:
                        run.text = f"[图片插入失败: {e}]"

    # --- 最终的目录处理 ---
    # 1. 从头开始，以编程方式添加一个全新的目录
    _add_toc(doc.docx)

    if not os.path.exists(GENERATED_DOCS_PATH):
        os.makedirs(GENERATED_DOCS_PATH)
        
    output_filename = "report_generated.docx"
    output_path = os.path.join(GENERATED_DOCS_PATH, output_filename)
    
    doc.save(output_path)

    try:
        _enable_auto_update_fields(output_path)
    except Exception as e:
        print(f"--- WARNING: Failed to enable auto-update fields for {output_path}. Reason: {e} ---")

    return output_path